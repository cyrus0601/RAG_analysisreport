"""
Description: This script generates a report based on the analysis data provided in a CSV file.
"""
# %%
import pandas as pd
import plotly.io as pio
import os

from docx import Document
from docx.shared import Inches
from plotly.subplots import make_subplots
from docx.shared import Inches
from docx2pdf import convert

# %%
def execute_code(code_str, idx):
        local_vars = {}
        image_paths = []
        try:
            # 是否需要 make_subplots
            if "make_subplots" in code_str:
                code_str = "from plotly.subplots import make_subplots\n" + code_str

            exec(code_str, {}, local_vars)
            figs = [fig for fig in local_vars.values() if hasattr(fig, 'write_image')]
            for fig_idx, fig in enumerate(figs):
                image_path = f'figure_{idx+1}_{fig_idx+1}.png'
                fig.write_image(image_path)
                image_paths.append(image_path)
        except Exception as e:
            print(f"Error executing code: {e}")
        return image_paths

def add_bold_paragraph(doc, text):
        p = doc.add_paragraph()
        segments = text.split('**')
        for i, segment in enumerate(segments):
            if i % 2 == 0:
                p.add_run(segment)
            else:
                p.add_run(segment).bold = True
                
def generate_report(csv_path, output_path):
    doc = Document()
    data = pd.read_csv(csv_path)

    analysis_blocks = data['Analysis'].tolist()
    code_blocks = data['Code'].tolist()

    # Define the titles
    titles = ["公司的業務之主要內容", "公司主要產品的重要用途", "公司的主要銷售地區", 
              "公司近兩年的銷售量及銷售值", "公司預測市場未來的供需及成長性", 
              "公司計畫開發的新產品", "公司短期和長期的業務發展計畫", "公司從業員工的基本情況"]
              
    doc.add_heading('公司分析報告', 0)

    for i, title in enumerate(titles):
        doc.add_heading(title, level=1)
        
        if i < len(analysis_blocks) and pd.notna(analysis_blocks[i]):
            lines = analysis_blocks[i].split("\n")
            for line in lines:
                if line.startswith("####"):
                    doc.add_heading(line.strip("####").strip(), level=4)
                elif line.startswith("###"):
                    doc.add_heading(line.strip("###").strip(), level=3)
                else:
                    add_bold_paragraph(doc, line)
        
        if i < len(code_blocks) and pd.notna(code_blocks[i]):
            image_paths = execute_code(code_blocks[i], i)
            if image_paths:
                for image_path in image_paths:
                    if os.path.exists(image_path):
                        doc.add_picture(image_path, width=Inches(6))
                    else:
                        doc.add_paragraph("未能生成圖片")
            
            for image_path in image_paths:
                if os.path.exists(image_path):
                    os.remove(image_path)

    doc.save(output_path)
    print(f"Document saved as {output_path}")

def convert_docx_to_pdf(docx_path, output_pdf_path=None):
    convert(docx_path, output_pdf_path)


#================================================================================================


# %%
csv_path = 'x_analysis.csv'
output_path = 'x分析報告.docx'
generate_report(csv_path, output_path)
convert_docx_to_pdf('x分析報告.docx', 'x分析報告.pdf')

# %%
csv_path = 'y_analysis.csv'
output_path = 'y分析報告.docx'
generate_report(csv_path, output_path)
convert_docx_to_pdf('y分析報告.docx', 'y分析報告.pdf')